import docx
#import json

class File_Manager:      
      def __init__(self):
          self.data_file = None
          self.image_file = None

          self.TEXT_STYLES = {
                  'Heading 1':'Title', 
                  'Heading 2':'Title', 
                  'Heading 3':'Title',
                  'Caption':'Summary',
                  'Body Text':'Body',
                  'Endnote':'Footer',
               }

      def set_data_document(self, data_file):
          self.data_file = data_file

      def set_image_file(self, image_file):
          self.image_file = image_file

      def get_image(self):
          return self.image

      def get_data_document(self):                
          document_object = docx.Document(self.data_file)

          for item in document_object.paragraphs:
              items = {"text_style_name":item.style.name, "text_body":item.text}              
              yield items
          
          '''for item in document_object.paragraphs:
              text_style = item.style.name
              block_text = self.TEXT_STYLES.get(text_style, 'Default Paragraph Style')
              print(f'{block_text}: {item.text}')
              print("\n")'''
          


      